import os
import re
import copy
import logging
import docx
import docx.oxml
import docx.oxml.ns
from docx.table import _Cell
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from typing import List, Dict, Any, Optional
from app.models import PaperConfig, Question

logger = logging.getLogger("app.generator")

def get_row_cell(table, row_idx: int, tc_idx: int) -> Optional[_Cell]:
    if row_idx >= len(table.rows):
        return None
    row = table.rows[row_idx]
    tc_list = row._tr.xpath('./w:tc')
    if tc_idx < len(tc_list):
        return _Cell(tc_list[tc_idx], table)
    return None

def set_cell_border_edge(cell, edge: str, val: str = 'single', sz: str = '4', space: str = '0', color: str = 'auto'):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(docx.oxml.ns.qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(tcBorders)
    
    for child in list(tcBorders):
        if child.tag.endswith(edge) or child.tag.endswith(f":{edge}") or child.tag == edge:
            tcBorders.remove(child)
    
    b_elem = parse_xml(f'<w:{edge} {nsdecls("w")} w:val="{val}" w:sz="{sz}" w:space="{space}" w:color="{color}"/>')
    tcBorders.append(b_elem)

def normalize_unit(unit_str: str) -> str:
    if not unit_str:
        return "Unit I"
    u = unit_str.strip().upper()
    if "III" in u or u == "UNIT 3" or u == "3":
        return "Unit III"
    if "II" in u or u == "UNIT 2" or u == "2":
        return "Unit II"
    if "IV" in u or u == "UNIT 4" or u == "4":
        return "Unit IV"
    if "V" in u or u == "UNIT 5" or u == "5":
        return "Unit V"
    if "I" in u or u == "UNIT 1" or u == "1":
        return "Unit I"
    return "Unit I"

def normalize_kl(kl_str: str) -> str:
    if not kl_str:
        return "K1"
    k = kl_str.strip().upper()
    if "K1" in k or "REMEMBER" in k: return "K1"
    if "K2" in k or "UNDERSTAND" in k: return "K2"
    if "K3" in k or "APPLY" in k or "APPLI" in k: return "K3"
    if "K4" in k or "ANALY" in k: return "K4"
    if "K5" in k or "EVALUAT" in k: return "K5"
    if "K6" in k or "CREAT" in k: return "K6"
    digits = re.findall(r'\d', k)
    if digits and 1 <= int(digits[0]) <= 6:
        return f"K{digits[0]}"
    return "K1"

def get_q_field(q, field: str, default: str = "") -> str:
    if not q:
        return default
    if isinstance(q, dict):
        val = q.get(field)
    else:
        val = getattr(q, field, None)
    return str(val) if val is not None else default

def set_cell_text_preserve_style(cell, text: str):
    """
    Clears the text in a cell while preserving its original cell borders and styles.
    Overwrites the text of the first run in the first paragraph, and removes other runs.
    """
    if len(cell.paragraphs) == 0:
        cell.add_paragraph()
    p = cell.paragraphs[0]
    
    if len(p.runs) > 0:
        first_run = p.runs[0]
        first_run.text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

def clean_degree_branch(deg_input: str) -> str:
    if not deg_input:
        return "B.E/CSE"
    cleaned = re.sub(r'\s*/\s*(?:[I|V|X]+|\d+)\s*$', '', deg_input.strip(), flags=re.IGNORECASE)
    cleaned = cleaned.replace("BE/BTECH", "B.E").replace("BE / BTECH", "B.E")
    cleaned = re.sub(r'\s*/\s*', '/', cleaned)
    return cleaned if cleaned else "B.E/CSE"

ROMAN_YEAR_SEM = {
    1: ("I", "I"), 2: ("I", "II"), 3: ("II", "III"), 4: ("II", "IV"),
    5: ("III", "V"), 6: ("III", "VI"), 7: ("IV", "VII"), 8: ("IV", "VIII")
}
ROMAN_MAP = {8: "VIII", 7: "VII", 6: "VI", 5: "V", 4: "IV", 3: "III", 2: "II", 1: "I"}

def format_year_sem(sem_input: str, alt_input: str = "") -> str:
    text = f"{sem_input or ''} {alt_input or ''}".upper()
    if not text.strip():
        return "II / III"

    m = re.search(r'\b([I|V|X]+)\s*/\s*([I|V|X]+)\b', text)
    if m:
        return f"{m.group(1)} / {m.group(2)}"

    for num in range(8, 0, -1):
        rom = ROMAN_MAP[num]
        if re.search(rf'\b({rom}|SEM\s*{num}|{num})\b', text):
            year_rom, sem_rom = ROMAN_YEAR_SEM[num]
            return f"{year_rom} / {sem_rom}"

    return " / "

def set_cell_bold_label_value(cell, label_bold: str, value_normal: str, font_size_pt: float = 11):
    if len(cell.paragraphs) == 0:
        cell.add_paragraph()
    p = cell.paragraphs[0]
    p.text = ""
    
    r1 = p.add_run(label_bold)
    r1.bold = True
    r1.font.name = "Times New Roman"
    r1.font.size = Pt(font_size_pt)
    
    if value_normal:
        r2 = p.add_run(f" {value_normal.strip()}")
        r2.bold = False
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(font_size_pt)

def replace_text_runs(doc, old_text: str, new_text: str):
    if not old_text:
        return
    if old_text == new_text:
        return
        
    def replace_in_paragraph(paragraph, old, new):
        if old not in paragraph.text:
            return
        
        search_start = 0
        while True:
            runs = paragraph.runs
            text = "".join(run.text for run in runs)
            start_idx = text.find(old, search_start)
            if start_idx == -1:
                break
                
            end_idx = start_idx + len(old)
            
            current_len = 0
            start_run_idx = -1
            start_run_offset = -1
            end_run_idx = -1
            end_run_offset = -1
            
            for i, run in enumerate(runs):
                run_len = len(run.text)
                if start_run_idx == -1 and current_len <= start_idx < current_len + run_len:
                    start_run_idx = i
                    start_run_offset = start_idx - current_len
                if current_len <= end_idx <= current_len + run_len:
                    end_run_idx = i
                    end_run_offset = end_idx - current_len
                    break
                current_len += run_len
                
            if start_run_idx != -1 and end_run_idx != -1:
                if start_run_idx == end_run_idx:
                    run = runs[start_run_idx]
                    run.text = run.text[:start_run_offset] + new + run.text[end_run_offset:]
                else:
                    start_run = runs[start_run_idx]
                    start_run.text = start_run.text[:start_run_offset] + new
                    
                    for r_idx in range(start_run_idx + 1, end_run_idx):
                        runs[r_idx].text = ""
                        
                    end_run = runs[end_run_idx]
                    end_run.text = end_run.text[end_run_offset:]
                
                search_start = start_idx + len(new)
            else:
                break

    for p in doc.paragraphs:
        replace_in_paragraph(p, old_text, new_text)
                
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, old_text, new_text)

def replace_set_placeholder(doc, target_set: str):
    if not target_set:
        return
    pattern = re.compile(r'\bSET\s*[\u2013\u2014\-–]?\s*(?:I{1,3}|IV|V|VI|VII|\d+)\b', re.IGNORECASE)
    for p in doc.paragraphs:
        matches = pattern.findall(p.text)
        for match_str in set(matches):
            replace_text_runs(doc, match_str, target_set)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    matches = pattern.findall(p.text)
                    for match_str in set(matches):
                        replace_text_runs(doc, match_str, target_set)

def replace_exam_title_placeholder(doc, exam_title: str):
    if not exam_title:
        return
    pattern = re.compile(r'CONTINUOUS ASSESSMENT TEST\s*[\u2013\u2014\-–]?\s*(?:I{1,3}|IV|V|VI|VII|\d+)', re.IGNORECASE)
    for p in doc.paragraphs:
        m = pattern.search(p.text)
        if m:
            replace_text_runs(doc, m.group(0), exam_title)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    m = pattern.search(p.text)
                    if m:
                        replace_text_runs(doc, m.group(0), exam_title)

def remove_table_rows(table, start_row_idx: int):
    """
    Safely removes rows from start_row_idx to the end of the table.
    """
    for row_idx in range(len(table.rows) - 1, start_row_idx - 1, -1):
        tr = table.rows[row_idx]._tr
        tr.getparent().remove(tr)

LEGACY_INSTITUTION_NAMES = {
    "NAME OF THE INSTITUTION:",
    "NAME OF THE INSTITUTION",
    "NAME OF THE INSTUTION:",
    "NAME OF THE INSTUTION",
    "JAYA ENGINEERING COLLEGE",
    "JAYA EDUCATIONAL TRUST",
    ""
}

def _generate_cat_paper(doc, config: PaperConfig, part_a: List[Question], part_b: List[List[Question]], part_c: List[Question]):
    # 1. Metadata Replacements
    inst_name = (config.institution_name or "").strip()
    if inst_name and not inst_name.upper().startswith("NAME OF THE INSTITUTION") and inst_name.upper() not in LEGACY_INSTITUTION_NAMES:
        replace_text_runs(doc, "NAME OF THE INSTITUTION:\t", inst_name.upper())
        replace_text_runs(doc, "NAME OF THE INSTITUTION :", inst_name.upper())
        replace_text_runs(doc, "NAME OF THE INSTITUTION:", inst_name.upper())

    if config.exam_type in ["CAT-3", "IAT-3"]:
        default_exam_title = "CONTINUOUS ASSESSMENT TEST - III"
    elif config.exam_type in ["CAT-2", "IAT-2"]:
        default_exam_title = "CONTINUOUS ASSESSMENT TEST - II"
    else:
        default_exam_title = "CONTINUOUS ASSESSMENT TEST - I"
    exam_title = config.exam_name or default_exam_title
    replace_exam_title_placeholder(doc, exam_title)
    
    if config.set:
        replace_set_placeholder(doc, config.set)
        
    reg_str = ""
    if config.regulation:
        reg_str = f"({config.regulation}-REGULATION)" if "REGULATION" not in config.regulation.upper() else config.regulation
        replace_text_runs(doc, "2021-REGULATION", reg_str)
    # Dynamically locate CAT Header Table and Course Details Table
    t_header = None
    t_course = None
    for t in doc.tables:
        txt = " ".join(c.text.strip() for row in t.rows for c in row.cells).upper()
        if ("SUB. CODE" in txt or "DEGREE / BRANCH" in txt or "DEGREE/BRANCH" in txt) and t_course is None:
            t_course = t
        elif ("NAME OF THE INSTITUTION" in txt or "CONTINUOUS ASSESSMENT TEST" in txt or "DATE/\nSESSION" in txt or "DATE / SESSION" in txt or "PAGES" in txt or "COPIES" in txt) and t_header is None:
            t_header = t

    # Fallback to index-based if not found dynamically
    if t_header is None:
        for t in doc.tables[:2]:
            if len(t.rows) >= 4:
                t_header = t
                break
    if t_course is None and len(doc.tables) > 1:
        for t in doc.tables[1:3]:
            if len(t.rows) in [3, 4] and len(t.rows[0].cells) <= 3:
                t_course = t
                break

    # Table 1: Header / Date & Session Table
    if t_header and len(t_header.rows) >= 4:
        # Check whether this template has 5 rows (CAT-3 format with Date & Session) or 4 rows (CAT-2 format)
        if len(t_header.rows) >= 5:
            # Row 1: Title & DATE/SESSION & Date value
            tcs1 = t_header.rows[1]._tr.xpath('./w:tc')
            if len(tcs1) >= 3:
                c_title = _Cell(tcs1[0], t_header)
                c_title.text = ""
                p = c_title.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(exam_title)
                r.bold = True
                r.underline = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(14)
                
                c_date = _Cell(tcs1[2], t_header)
                c_date.text = ""
                date_clean = (config.date or "").replace("_", "").strip()
                if date_clean:
                    p = c_date.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(date_clean)
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(12)

            # Row 2: Session value
            tcs2 = t_header.rows[2]._tr.xpath('./w:tc')
            if len(tcs2) >= 3:
                c_sess = _Cell(tcs2[2], t_header)
                c_sess.text = ""
                sess_clean = (config.session or "").replace("_", "").strip()
                if sess_clean:
                    p = c_sess.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(sess_clean)
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(12)

            # Row 3: Regulation & PAGES
            tcs3 = t_header.rows[3]._tr.xpath('./w:tc')
            if len(tcs3) >= 3:
                c_reg = _Cell(tcs3[0], t_header)
                c_reg.text = ""
                p = c_reg.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                reg_text = reg_str if reg_str else "(2021-REGULATION)"
                if not reg_text.startswith("("):
                    reg_text = f"({reg_text})"
                r = p.add_run(reg_text)
                r.font.name = "Times New Roman"
                r.font.size = Pt(14)

            # Row 4: Semester & COPIES
            tcs4 = t_header.rows[4]._tr.xpath('./w:tc')
            if len(tcs4) >= 3:
                c_sem = _Cell(tcs4[0], t_header)
                c_sem.text = ""
                p = c_sem.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sem_text = config.semester.strip() if (config.semester and not re.match(r'^[IVX\s/]+$', config.semester.strip(), re.I)) else "ODD SEMESTER 2026-27"
                r = p.add_run(sem_text)
                r.font.name = "Times New Roman"
                r.font.size = Pt(14)

        elif len(t_header.rows) == 4:
            # 4-row layout (CAT-2)
            # Row 1: Title & DATE & Date value
            tcs1 = t_header.rows[1]._tr.xpath('./w:tc')
            if len(tcs1) >= 3:
                c_title = _Cell(tcs1[0], t_header)
                c_title.text = ""
                p = c_title.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(exam_title)
                r.bold = True
                r.underline = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(14)

                c_date = _Cell(tcs1[2], t_header)
                c_date.text = ""
                date_clean = (config.date or "").replace("_", "").strip()
                if date_clean:
                    p = c_date.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(date_clean)
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(12)

            # Row 2: Regulation & PAGES
            tcs2 = t_header.rows[2]._tr.xpath('./w:tc')
            if len(tcs2) >= 3:
                c_reg = _Cell(tcs2[0], t_header)
                c_reg.text = ""
                p = c_reg.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                reg_text = reg_str if reg_str else "(2021-REGULATION)"
                if not reg_text.startswith("("):
                    reg_text = f"({reg_text})"
                r = p.add_run(reg_text)
                r.font.name = "Times New Roman"
                r.font.size = Pt(14)

            # Row 3: Semester & COPIES
            tcs3 = t_header.rows[3]._tr.xpath('./w:tc')
            if len(tcs3) >= 3:
                c_sem = _Cell(tcs3[0], t_header)
                c_sem.text = ""
                p = c_sem.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sem_text = config.semester.strip() if (config.semester and not re.match(r'^[IVX\s/]+$', config.semester.strip(), re.I)) else "ODD SEMESTER 2026-27"
                r = p.add_run(sem_text)
                r.font.name = "Times New Roman"
                r.font.size = Pt(14)

    # Table 2: Course & Exam Info Box
    if t_course:
        sub_code = (config.subject_code or "").strip()
        sub_name = (config.subject_name or "").strip()
        if sub_code and sub_name:
            sub_val = f"{sub_code} – {sub_name}"
        elif sub_code:
            sub_val = sub_code
        elif sub_name:
            sub_val = sub_name
        else:
            sub_val = ""

        if len(t_course.rows) > 0 and len(t_course.rows[0].cells) > 0:
            set_cell_bold_label_value(t_course.rows[0].cells[0], "Sub. Code / Sub. Name  :", sub_val)
            if len(t_course.rows[0].cells) > 1:
                set_cell_bold_label_value(t_course.rows[0].cells[1], "Sub. Code / Sub. Name  :", sub_val)
            
        deg_branch = clean_degree_branch(config.degree_branch_sem)
        if len(t_course.rows) > 1 and len(t_course.rows[1].cells) > 0:
            set_cell_bold_label_value(t_course.rows[1].cells[0], "Degree / Branch:", deg_branch)
            
        sem_val = format_year_sem(config.degree_branch_sem, config.semester)
        if len(t_course.rows) > 1 and len(t_course.rows[1].cells) > 1:
            set_cell_bold_label_value(t_course.rows[1].cells[1], "Year / Semester:", sem_val)
            
        time_val = (config.time or "").strip() or "90 Minutes"
        if len(t_course.rows) > 2 and len(t_course.rows[2].cells) > 0:
            set_cell_bold_label_value(t_course.rows[2].cells[0], "Time:", time_val)
            
        marks_val = str(config.max_marks) if config.max_marks else "50"
        if len(t_course.rows) > 2 and len(t_course.rows[2].cells) > 1:
            set_cell_bold_label_value(t_course.rows[2].cells[1], "Maximum Marks:", marks_val)

    # 2. Part A (Table 3)
    # 2. Find Question Tables (Part A, Part B, Part C) dynamically
    q_tables = []
    for t in doc.tables:
        if len(t.rows) > 0 and len(t.rows[0].cells) in [4, 5, 7]:
            header = " ".join(c.text.strip() for c in t.rows[0].cells).upper()
            if "Q.NO" in header and ("QUESTION" in header or "QUESTIONS" in header):
                q_tables.append(t)

    t_part_a = q_tables[0] if len(q_tables) > 0 else (doc.tables[3] if len(doc.tables) > 3 else None)
    t_part_b = q_tables[1] if len(q_tables) > 1 else (doc.tables[4] if len(doc.tables) > 4 else None)
    t_part_c = q_tables[2] if len(q_tables) > 2 else (doc.tables[5] if len(doc.tables) > 5 else None)

    # 2. Part A
    if t_part_a:
        for idx, q in enumerate(part_a[:5]):
            row_idx = 1 + idx
            if row_idx < len(t_part_a.rows):
                if len(t_part_a.rows[row_idx].cells) > 1:
                    set_cell_text_preserve_style(t_part_a.rows[row_idx].cells[1], get_q_field(q, "text"))
                if len(t_part_a.rows[row_idx].cells) > 2:
                    set_cell_text_preserve_style(t_part_a.rows[row_idx].cells[2], get_q_field(q, "kl"))
                if len(t_part_a.rows[row_idx].cells) > 3:
                    set_cell_text_preserve_style(t_part_a.rows[row_idx].cells[3], get_q_field(q, "co"))

    # 3. Part B & Part C based on template table structure
    is_2025_cat_layout = False
    if t_part_b and len(t_part_b.rows) > 0 and len(t_part_b.rows[0].cells) == 4:
        is_2025_cat_layout = True

    if is_2025_cat_layout and t_part_b:
        # 2025 Regulation: Part B (Q6..Q10 short questions)
        flat_b = []
        for item in part_b:
            if isinstance(item, list):
                flat_b.extend([q for q in item if q])
            elif item:
                flat_b.append(item)

        for idx, q in enumerate(flat_b[:5]):
            row_idx = 1 + idx
            if row_idx < len(t_part_b.rows):
                if len(t_part_b.rows[row_idx].cells) > 1:
                    set_cell_text_preserve_style(t_part_b.rows[row_idx].cells[1], get_q_field(q, "text"))
                if len(t_part_b.rows[row_idx].cells) > 2:
                    set_cell_text_preserve_style(t_part_b.rows[row_idx].cells[2], get_q_field(q, "kl"))
                if len(t_part_b.rows[row_idx].cells) > 3:
                    set_cell_text_preserve_style(t_part_b.rows[row_idx].cells[3], get_q_field(q, "co"))

        # 2025 Regulation: Part C (Q11a/11b, Q12a/12b, Q13a/13b)
        if t_part_c:
            flat_c = []
            for item in part_c:
                if isinstance(item, list):
                    flat_c.extend([q for q in item if q])
                elif item:
                    flat_c.append(item)

            row_indices = [1, 3, 4, 6, 7, 9] # Q11a, Q11b, Q12a, Q12b, Q13a, Q13b
            for idx, q in enumerate(flat_c[:6]):
                if idx < len(row_indices):
                    r_idx = row_indices[idx]
                    if r_idx < len(t_part_c.rows):
                        r = t_part_c.rows[r_idx]
                        unique_cells = []
                        for cell in r.cells:
                            if not any(uc._tc == cell._tc for uc in unique_cells):
                                unique_cells.append(cell)
                        if len(unique_cells) >= 5:
                            set_cell_text_preserve_style(unique_cells[2], get_q_field(q, "text"))
                            set_cell_text_preserve_style(unique_cells[3], get_q_field(q, "kl"))
                            set_cell_text_preserve_style(unique_cells[4], get_q_field(q, "co"))
                        else:
                            n_cells = len(r.cells)
                            if n_cells >= 4:
                                set_cell_text_preserve_style(r.cells[3], get_q_field(q, "text"))
                            col_kl = n_cells - 2 if n_cells >= 6 else 2
                            col_co = n_cells - 1 if n_cells >= 6 else 3
                            if col_kl < n_cells:
                                set_cell_text_preserve_style(r.cells[col_kl], get_q_field(q, "kl"))
                            if col_co < n_cells:
                                set_cell_text_preserve_style(r.cells[col_co], get_q_field(q, "co"))

    else:
        # 2021 Regulation: Part B (Either-Or pairs Q6a/b, Q7a/b)
        if t_part_b:
            pair_rows = [(1, 3), (4, 6)]
            for idx, pair in enumerate(part_b[:2]):
                if idx < len(pair_rows) and isinstance(pair, (list, tuple)):
                    row_a_idx, row_b_idx = pair_rows[idx]
                    if row_a_idx < len(t_part_b.rows) and len(pair) > 0 and pair[0]:
                        q_a = pair[0]
                        if len(t_part_b.rows[row_a_idx].cells) > 2:
                            set_cell_text_preserve_style(t_part_b.rows[row_a_idx].cells[2], get_q_field(q_a, "text"))
                        if len(t_part_b.rows[row_a_idx].cells) > 3:
                            set_cell_text_preserve_style(t_part_b.rows[row_a_idx].cells[3], get_q_field(q_a, "kl"))
                        if len(t_part_b.rows[row_a_idx].cells) > 4:
                            set_cell_text_preserve_style(t_part_b.rows[row_a_idx].cells[4], get_q_field(q_a, "co"))
                    if row_b_idx < len(t_part_b.rows) and len(pair) > 1 and pair[1]:
                        q_b = pair[1]
                        if len(t_part_b.rows[row_b_idx].cells) > 2:
                            set_cell_text_preserve_style(t_part_b.rows[row_b_idx].cells[2], get_q_field(q_b, "text"))
                        if len(t_part_b.rows[row_b_idx].cells) > 3:
                            set_cell_text_preserve_style(t_part_b.rows[row_b_idx].cells[3], get_q_field(q_b, "kl"))
                        if len(t_part_b.rows[row_b_idx].cells) > 4:
                            set_cell_text_preserve_style(t_part_b.rows[row_b_idx].cells[4], get_q_field(q_b, "co"))

        # 2021 Regulation: Part C (Either-Or pair Q8a/b in Table Part C)
        if t_part_c:
            flat_c = []
            for item in part_c:
                if isinstance(item, (list, tuple)):
                    flat_c.extend([q for q in item if q])
                elif item:
                    flat_c.append(item)

            if len(flat_c) >= 1 and len(t_part_c.rows) > 1:
                q_a = flat_c[0]
                if len(t_part_c.rows[1].cells) > 2:
                    set_cell_text_preserve_style(t_part_c.rows[1].cells[2], get_q_field(q_a, "text"))
                if len(t_part_c.rows[1].cells) > 3:
                    set_cell_text_preserve_style(t_part_c.rows[1].cells[3], get_q_field(q_a, "kl"))
                if len(t_part_c.rows[1].cells) > 4:
                    set_cell_text_preserve_style(t_part_c.rows[1].cells[4], get_q_field(q_a, "co"))
            if len(flat_c) >= 2 and len(t_part_c.rows) > 3:
                q_b = flat_c[1]
                if len(t_part_c.rows[3].cells) > 2:
                    set_cell_text_preserve_style(t_part_c.rows[3].cells[2], get_q_field(q_b, "text"))
                if len(t_part_c.rows[3].cells) > 3:
                    set_cell_text_preserve_style(t_part_c.rows[3].cells[3], get_q_field(q_b, "kl"))
                if len(t_part_c.rows[3].cells) > 4:
                    set_cell_text_preserve_style(t_part_c.rows[3].cells[4], get_q_field(q_b, "co"))


    # 4. Table of Specifications (TOS) for CAT
    is_cat3 = config.exam_type in ["CAT-3", "IAT-3"]
    is_cat2 = config.exam_type in ["CAT-2", "IAT-2"]
    is_2025 = (config.regulation == "2025") if config.regulation else is_2025_cat_layout
    if is_cat3:
        target_units = ["Unit IV", "Unit V"]
        unit_labels = ["IV", "V"]
    elif is_cat2 and not is_2025:
        target_units = ["Unit II", "Unit III"]
        unit_labels = ["II", "III"]
    elif is_cat2:
        target_units = ["Unit III", "Unit IV"]
        unit_labels = ["III", "IV"]
    else:
        target_units = ["Unit I", "Unit II"]
        unit_labels = ["I", "II"]

    kls_map = {"K1": 0, "K2": 1, "K3": 2, "K4": 3, "K5": 4, "K6": 5}
    tos_counts = [[0 for _ in range(6)] for _ in range(2)]
    tos_marks = [[0 for _ in range(6)] for _ in range(2)]

    is_2025 = (config.regulation == "2025") if config.regulation else is_2025_cat_layout
    part_a_mark = 1 if is_2025 else 2
    part_b_mark = 3 if is_2025 else 13
    part_c_mark = 10 if is_2025 else 14

    questions_with_marks = []
    for q in part_a:
        if q:
            questions_with_marks.append((q, part_a_mark))

    for item in part_b:
        if isinstance(item, list):
            for q in item:
                if q:
                    questions_with_marks.append((q, part_b_mark))
        elif item:
            questions_with_marks.append((item, part_b_mark))

    for item in part_c:
        if isinstance(item, list):
            for q in item:
                if q:
                    questions_with_marks.append((q, part_c_mark))
        elif item:
            questions_with_marks.append((item, part_c_mark))

    for q, section_mark in questions_with_marks:
        u_norm = normalize_unit(get_q_field(q, "unit", "Unit I"))
        if u_norm == target_units[0]:
            u_idx = 0
        elif u_norm == target_units[1]:
            u_idx = 1
        else:
            u_idx = 0
        
        kl_key = normalize_kl(get_q_field(q, "kl", "K1"))
        kl_idx = kls_map.get(kl_key, 0)

        tos_counts[u_idx][kl_idx] += 1
        tos_marks[u_idx][kl_idx] += section_mark

    # Locate TOS tables dynamically by checking table content
    t6 = None  # Question-wise TOS
    t7 = None  # Marks-wise TOS
    for t in doc.tables:
        if len(t.rows) > 0 and len(t.rows[0].cells) >= 8:
            header_text = " ".join(c.text for row in t.rows[:2] for c in row.cells).upper()
            if "SYLLABUS" in header_text or "NO. OF QUESTIONS" in header_text or "MARKS" in header_text:
                if t6 is None:
                    t6 = t
                elif t7 is None:
                    t7 = t

    # Fallback to index-based if not found
    if t6 is None and len(doc.tables) > 6:
        t6 = doc.tables[6]
    if t7 is None and len(doc.tables) > 7:
        t7 = doc.tables[7]

    # Table 6 (Question-wise TOS)
    if t6 is not None and len(t6.rows) >= 4:
        unit_row_start = 2 if len(t6.rows) == 5 else (3 if len(t6.rows) > 5 else 2)
        total_row_idx = len(t6.rows) - 1

        for u_idx in range(2):
            row_idx = unit_row_start + u_idx
            if row_idx < total_row_idx and len(t6.rows[row_idx].cells) >= 8:
                set_cell_text_preserve_style(t6.rows[row_idx].cells[0], unit_labels[u_idx])
                row_sum = 0
                for k_idx in range(6):
                    val = tos_counts[u_idx][k_idx]
                    set_cell_text_preserve_style(t6.rows[row_idx].cells[1 + k_idx], str(val) if val > 0 else "")
                    row_sum += val
                set_cell_text_preserve_style(t6.rows[row_idx].cells[7], str(row_sum))

        # Total row in Table 6
        if total_row_idx < len(t6.rows) and len(t6.rows[total_row_idx].cells) >= 8:
            set_cell_text_preserve_style(t6.rows[total_row_idx].cells[0], "Total")
            for k_idx in range(6):
                col_sum = sum(tos_counts[u_idx][k_idx] for u_idx in range(2))
                set_cell_text_preserve_style(t6.rows[total_row_idx].cells[1 + k_idx], str(col_sum) if col_sum > 0 else "0")
            grand_total = sum(sum(r) for r in tos_counts)
            set_cell_text_preserve_style(t6.rows[total_row_idx].cells[7], str(grand_total))

    # Table 7 (Marks-wise TOS)
    if t7 is not None and len(t7.rows) >= 4:
        unit_row_start = 2 if len(t7.rows) == 5 else (3 if len(t7.rows) > 5 else 2)
        total_row_idx = len(t7.rows) - 1

        for u_idx in range(2):
            row_idx = unit_row_start + u_idx
            if row_idx < total_row_idx and len(t7.rows[row_idx].cells) >= 8:
                set_cell_text_preserve_style(t7.rows[row_idx].cells[0], unit_labels[u_idx])
                row_sum = 0
                for k_idx in range(6):
                    val = tos_marks[u_idx][k_idx]
                    set_cell_text_preserve_style(t7.rows[row_idx].cells[1 + k_idx], str(val) if val > 0 else "")
                    row_sum += val
                set_cell_text_preserve_style(t7.rows[row_idx].cells[7], str(row_sum))

        # Total row in Table 7
        if total_row_idx < len(t7.rows) and len(t7.rows[total_row_idx].cells) >= 8:
            set_cell_text_preserve_style(t7.rows[total_row_idx].cells[0], "Total")
            for k_idx in range(6):
                col_sum = sum(tos_marks[u_idx][k_idx] for u_idx in range(2))
                set_cell_text_preserve_style(t7.rows[total_row_idx].cells[1 + k_idx], str(col_sum) if col_sum > 0 else "0")
            grand_total = sum(sum(r) for r in tos_marks)
            set_cell_text_preserve_style(t7.rows[total_row_idx].cells[7], str(grand_total))


def _generate_model_paper(doc, config: PaperConfig, part_a: List[Question], part_b: List[List[Question]], part_c: List[Question]):
    # 1. Replace metadata placeholders
    if config.institution_name and config.institution_name.strip().upper() not in LEGACY_INSTITUTION_NAMES:
        replace_text_runs(doc, "NAME OF THE INSTITUTION :", config.institution_name.upper())
        replace_text_runs(doc, "NAME OF THE INSTITUTION:", config.institution_name.upper())

    sub_code = (config.subject_code or "").strip()
    sub_name = (config.subject_name or "").strip()
    deg_branch = (config.degree_branch_sem or "").strip()
    sem_val = (config.semester or "").strip()

    for p in doc.paragraphs:
        p_text = p.text
        if "Sub. Code" in p_text or "Sub.Code" in p_text or "Sub.Name" in p_text:
            if sub_code and sub_name:
                p.text = f"Sub. Code/Sub.Name: {sub_code}/ {sub_name}"
            elif sub_code:
                p.text = f"Sub. Code/Sub.Name: {sub_code}"
            elif sub_name:
                p.text = f"Sub. Code/Sub.Name: {sub_name}"
        elif "Degree/Branch/Sem" in p_text or ("Degree" in p_text and "Branch" in p_text):
            parts = []
            if deg_branch:
                parts.append(deg_branch)
            if sem_val and sem_val not in deg_branch:
                parts.append(sem_val)
            full_deg = " / ".join(parts) if parts else "BE/BTECH"
            p.text = f"Degree/Branch/Sem: {full_deg}"

    if config.subject_code:
        replace_text_runs(doc, "OCS353", config.subject_code)
    if config.subject_name:
        replace_text_runs(doc, "Data Science fundamentals", config.subject_name)
    if config.regulation:
        replace_text_runs(doc, "2021-Regulation", config.regulation)
    if config.semester:
        replace_text_runs(doc, "ODD SEMESTER-2025-26", config.semester)
    if config.time:
        replace_text_runs(doc, "3 Hours", config.time)
    if config.max_marks:
        replace_text_runs(doc, "Maximum Marks: 100", f"Maximum Marks: {config.max_marks}")
    if config.set:
        replace_set_placeholder(doc, config.set)
    if config.exam_name:
        replace_text_runs(doc, "MODEL EXAMINATION", config.exam_name)
    if config.date:
        replace_text_runs(doc, "Date:", f"Date: {config.date}")
        
    # 2. Populate Part A (Table 1)
    t1 = doc.tables[1]
    for idx, q in enumerate(part_a):
        row_idx = 1 + idx
        if row_idx < len(t1.rows):
            set_cell_text_preserve_style(t1.rows[row_idx].cells[1], get_q_field(q, "text"))
            set_cell_text_preserve_style(t1.rows[row_idx].cells[2], get_q_field(q, "kl"))
            set_cell_text_preserve_style(t1.rows[row_idx].cells[3], get_q_field(q, "co"))

    # If Part A has fewer questions (e.g. 5 for CAT-1/CAT-2), trim extra rows
    if len(part_a) < 10 and len(t1.rows) > (1 + len(part_a)):
        remove_table_rows(t1, 1 + len(part_a))
            
    # 3. Populate Part B (Table 2)
    t2 = doc.tables[2]
    for idx, pair in enumerate(part_b):
        row_a_idx = 1 + idx * 3
        row_b_idx = 3 + idx * 3
        
        if row_a_idx < len(t2.rows) and len(pair) > 0:
            q_a = pair[0]
            set_cell_text_preserve_style(t2.rows[row_a_idx].cells[2], get_q_field(q_a, "text"))
            set_cell_text_preserve_style(t2.rows[row_a_idx].cells[3], get_q_field(q_a, "kl"))
            set_cell_text_preserve_style(t2.rows[row_a_idx].cells[4], get_q_field(q_a, "co"))
            
        if row_b_idx < len(t2.rows) and len(pair) > 1:
            q_b = pair[1]
            set_cell_text_preserve_style(t2.rows[row_b_idx].cells[2], get_q_field(q_b, "text"))
            set_cell_text_preserve_style(t2.rows[row_b_idx].cells[3], get_q_field(q_b, "kl"))
            set_cell_text_preserve_style(t2.rows[row_b_idx].cells[4], get_q_field(q_b, "co"))

    # If Part B has fewer pairs (e.g. 2 for CAT-1/CAT-2), trim extra rows
    if len(part_b) < 5 and len(t2.rows) > (1 + len(part_b) * 3):
        remove_table_rows(t2, 1 + len(part_b) * 3)

    # 4. Populate Part C (Table 3)
    t3 = doc.tables[3]
    if len(part_c) >= 2:
        set_cell_text_preserve_style(t3.rows[1].cells[2], get_q_field(part_c[0], "text"))
        set_cell_text_preserve_style(t3.rows[1].cells[3], get_q_field(part_c[0], "kl"))
        set_cell_text_preserve_style(t3.rows[1].cells[4], get_q_field(part_c[0], "co"))
        
        set_cell_text_preserve_style(t3.rows[3].cells[2], get_q_field(part_c[1], "text"))
        set_cell_text_preserve_style(t3.rows[3].cells[3], get_q_field(part_c[1], "kl"))
        set_cell_text_preserve_style(t3.rows[3].cells[4], get_q_field(part_c[1], "co"))

    # 5. Compute Table of Specifications (TOS)
    units_map = {"Unit I": 0, "Unit II": 1, "Unit III": 2, "Unit IV": 3, "Unit V": 4}
    kls_map = {"K1": 0, "K2": 1, "K3": 2, "K4": 3, "K5": 4, "K6": 5}
    
    tos_counts = [[0 for _ in range(6)] for _ in range(5)]
    tos_marks = [[0 for _ in range(6)] for _ in range(5)]
    
    is_2025 = (config.regulation == "2025") if config.regulation else False
    part_a_mark = 1 if is_2025 else 2
    part_b_mark = 3 if is_2025 else 13
    part_c_mark = 10 if is_2025 else 15

    questions_with_marks = []
    for q in part_a:
        if q:
            questions_with_marks.append((q, part_a_mark))
            
    for pair in part_b:
        if isinstance(pair, list):
            for q in pair:
                if q:
                    questions_with_marks.append((q, part_b_mark))
        elif pair:
            questions_with_marks.append((pair, part_b_mark))
            
    for item in part_c:
        if isinstance(item, list):
            for q in item:
                if q:
                    questions_with_marks.append((q, part_c_mark))
        elif item:
            questions_with_marks.append((item, part_c_mark))
            
    for q, section_mark in questions_with_marks:
        u_norm = normalize_unit(get_q_field(q, "unit", "Unit I"))
        unit_idx = units_map.get(u_norm)
        kl_key = normalize_kl(get_q_field(q, "kl", "K1"))
        kl_idx = kls_map.get(kl_key)
        
        if unit_idx is not None and kl_idx is not None:
            tos_counts[unit_idx][kl_idx] += 1
            tos_marks[unit_idx][kl_idx] += section_mark

    # 6. Populate Table 4 (Question-Wise TOS)
    t4 = doc.tables[4]
    for u_idx in range(5):
        row_idx = 2 + u_idx
        row_sum = 0
        for k_idx in range(6):
            val = tos_counts[u_idx][k_idx]
            val_str = str(val) if val > 0 else ""
            set_cell_text_preserve_style(t4.rows[row_idx].cells[1 + k_idx], val_str)
            row_sum += val
        set_cell_text_preserve_style(t4.rows[row_idx].cells[7], str(row_sum))
        
    for k_idx in range(6):
        col_sum = sum(tos_counts[u_idx][k_idx] for u_idx in range(5))
        set_cell_text_preserve_style(t4.rows[7].cells[1 + k_idx], str(col_sum) if col_sum > 0 else "0")
    grand_total_t4 = sum(sum(row) for row in tos_counts)
    set_cell_text_preserve_style(t4.rows[7].cells[7], str(grand_total_t4))

    # 7. Populate Table 5 (Marks-Wise TOS)
    t5 = doc.tables[5]
    for u_idx in range(5):
        row_idx = 2 + u_idx
        row_sum = 0
        for k_idx in range(6):
            val = tos_marks[u_idx][k_idx]
            val_str = str(val) if val > 0 else ""
            set_cell_text_preserve_style(t5.rows[row_idx].cells[1 + k_idx], val_str)
            row_sum += val
        set_cell_text_preserve_style(t5.rows[row_idx].cells[7], str(row_sum))
        
    for k_idx in range(6):
        col_sum = sum(tos_marks[u_idx][k_idx] for u_idx in range(5))
        set_cell_text_preserve_style(t5.rows[7].cells[1 + k_idx], str(col_sum) if col_sum > 0 else "0")
    grand_total_t5 = sum(sum(row) for row in tos_marks)
    set_cell_text_preserve_style(t5.rows[7].cells[7], str(grand_total_t5))


def generate_question_paper(
    template_path: str,
    output_path: Any,
    config: PaperConfig,
    part_a: List[Question],
    part_b: List[List[Question]],
    part_c: List[Question]
):
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template Word document not found at '{template_path}'")

    doc = docx.Document(template_path)
    
    if len(doc.tables) < 6:
        raise ValueError(f"Template document structure invalid. Expected at least 6 tables, found {len(doc.tables)}.")

    logger.info(f"Generating question paper for {config.subject_code} ({config.set}) - Exam Type: {config.exam_type}")

    is_cat_template = len(doc.tables) >= 9 or "cat" in os.path.basename(template_path).lower()
    
    if is_cat_template:
        _generate_cat_paper(doc, config, part_a, part_b, part_c)
    else:
        _generate_model_paper(doc, config, part_a, part_b, part_c)
        
    doc.save(output_path)
    logger.info(f"Successfully generated question paper document at: {output_path}")
