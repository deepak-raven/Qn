import io
import os
import re
import shutil
import logging
import platform
import subprocess
import docx
from docx.text.paragraph import Paragraph
from docx.table import Table
import pdfplumber
from typing import List, Dict, Any, Optional

logger = logging.getLogger("app.parser")


def parse_unit_from_text(text: str) -> Optional[str]:
    if not text:
        return None
    for line in text.split('\n'):
        line_clean = re.sub(r'\s+', ' ', line.upper()).strip()
        # Match UNIT I, UNIT-1, UNIT: 01, MODULE 1, CHAPTER 1, etc.
        m = re.search(r'\b(?:UNIT|MODULE|CHAPTER|TOPIC)\s*[-–:]?\s*([IVX\d]+|\bONE\b|\bTWO\b|\bTHREE\b|\bFOUR\b|\bFIVE\b)\b', line_clean)
        if m:
            u = m.group(1).strip()
            mapping = {
                'I': 'Unit I', '1': 'Unit I', '01': 'Unit I', 'ONE': 'Unit I',
                'II': 'Unit II', '2': 'Unit II', '02': 'Unit II', 'TWO': 'Unit II',
                'III': 'Unit III', '3': 'Unit III', '03': 'Unit III', 'THREE': 'Unit III',
                'IV': 'Unit IV', '4': 'Unit IV', '04': 'Unit IV', 'FOUR': 'Unit IV',
                'V': 'Unit V', '5': 'Unit V', '05': 'Unit V', 'FIVE': 'Unit V',
            }
            if u in mapping:
                return mapping[u]
            if u.isdigit():
                num = int(u)
                roman_map = {1: 'Unit I', 2: 'Unit II', 3: 'Unit III', 4: 'Unit IV', 5: 'Unit V'}
                return roman_map.get(num, f"Unit {num}")
    return None


def parse_part_marks_from_text(text: str):
    if not text:
        return None, None
    for line in text.split('\n'):
        line_clean = re.sub(r'\s+', ' ', line.lower()).strip()
        if len(line_clean) <= 150:
            # Check Part A / 1 or 2 Marks
            if (re.search(r'\b(one|1)\s*marks?\b', line_clean) or 
                re.search(r'\b(two|2)\s*marks?\b', line_clean) or 
                re.search(r'\bpart\s*[-–:]?\s*a\b', line_clean)):
                m_val = 1 if ('one' in line_clean or '1' in line_clean) else 2
                return 'A', m_val

            # Check Part B / 3, 13, 16 Marks
            if (re.search(r'\b(three|3)\s*marks?\b', line_clean) or 
                re.search(r'\b(thirteen|13)\s*marks?\b', line_clean) or 
                re.search(r'\b(sixteen|16)\s*marks?\b', line_clean) or 
                re.search(r'\bpart\s*[-–:]?\s*b\b', line_clean)):
                if 'three' in line_clean or '3' in line_clean:
                    m_val = 3
                elif 'sixteen' in line_clean or '16' in line_clean:
                    m_val = 16
                else:
                    m_val = 13
                return 'B', m_val

            # Check Part C / 10, 12, 14, 15, 16 Marks
            if (re.search(r'\b(ten|10)\s*marks?\b', line_clean) or 
                re.search(r'\b(twelve|12)\s*marks?\b', line_clean) or 
                re.search(r'\b(fourteen|14)\s*marks?\b', line_clean) or 
                re.search(r'\b(fifteen|15)\s*marks?\b', line_clean) or 
                re.search(r'\bpart\s*[-–:]?\s*c\b', line_clean)):
                if 'ten' in line_clean or '10' in line_clean:
                    m_val = 10
                elif 'twelve' in line_clean or '12' in line_clean:
                    m_val = 12
                elif 'fourteen' in line_clean or '14' in line_clean:
                    m_val = 14
                elif 'sixteen' in line_clean or '16' in line_clean:
                    m_val = 16
                else:
                    m_val = 15
                return 'C', m_val
    return None, None


def infer_co_from_unit(unit_str: Optional[str]) -> str:
    if not unit_str:
        return "CO1"
    u = unit_str.upper()
    if "I" in u and "III" not in u and "II" not in u and "IV" not in u and "VIII" not in u and "VII" not in u:
        return "CO1"
    if "II" in u and "III" not in u and "VII" not in u and "VIII" not in u:
        return "CO2"
    if "III" in u and "VIII" not in u:
        return "CO3"
    if "IV" in u:
        return "CO4"
    if "V" in u:
        return "CO5"
    if "1" in u: return "CO1"
    if "2" in u: return "CO2"
    if "3" in u: return "CO3"
    if "4" in u: return "CO4"
    if "5" in u: return "CO5"
    return "CO1"


def infer_bloom_from_text(question_text: str, part: str = 'A') -> str:
    if not question_text:
        return "K1" if part == 'A' else "K2"
    q_low = question_text.lower().strip()
    
    # Bloom Level 1 - Remember
    if re.search(r'\b(define|state|list|name|what is|what are|mention|recall|identify|give the term|write down|who|when|where|label|tabulate)\b', q_low):
        return "K1"
    # Bloom Level 2 - Understand
    if re.search(r'\b(explain|describe|discuss|differentiate|distinguish|compare|classify|illustrate|summarize|outline|contrast|why|interpret|paraphrase|indicate)\b', q_low):
        return "K2"
    # Bloom Level 3 - Apply
    if re.search(r'\b(apply|calculate|solve|demonstrate|construct|implement|compute|determine|show how|use|execute|modify|show that|find)\b', q_low):
        return "K3"
    # Bloom Level 4 - Analyze
    if re.search(r'\b(analyze|analyse|examine|categorize|breakdown|differentiate between|investigate|inspect|test|probe)\b', q_low):
        return "K4"
    # Bloom Level 5 - Evaluate
    if re.search(r'\b(evaluate|assess|justify|criticize|judge|defend|prioritize|rate|appraise|recommend)\b', q_low):
        return "K5"
    # Bloom Level 6 - Create
    if re.search(r'\b(design|formulate|develop|create|compose|plan|propose|synthesize|generate|devise)\b', q_low):
        return "K6"
        
    return "K1" if part == 'A' else "K2"


def extract_inline_metadata(raw_text: str) -> tuple:
    """Extracts inline (K1-K6), (CO1-CO6), and marks [2] from question body if present."""
    if not raw_text:
        return "", "", "", None
        
    text = raw_text.strip()
    kl = ""
    co = ""
    marks = None
    
    # Extract KL e.g. (K1), [K2], (BT-2), (Level 3), [Bloom Level 2]
    kl_match = re.search(r'[\(\[\{]\s*(?:K|KL|BT|LEVEL|BLOOM LEVEL|BLOOM)?\s*[-–:]?\s*([1-6]|K[1-6]|REMEMBER|UNDERSTAND|APPLY|ANALYZE|EVALUATE|CREATE)\s*[\)\]\}]', text, re.IGNORECASE)
    if kl_match:
        val = kl_match.group(1).upper()
        if val.startswith('K') and len(val) == 2:
            kl = val
        elif val.isdigit():
            kl = f"K{val}"
        elif "REMEMBER" in val: kl = "K1"
        elif "UNDERSTAND" in val: kl = "K2"
        elif "APPLY" in val: kl = "K3"
        elif "ANALY" in val: kl = "K4"
        elif "EVALUAT" in val: kl = "K5"
        elif "CREAT" in val: kl = "K6"
        text = text[:kl_match.start()] + text[kl_match.end():]
        
    # Extract CO e.g. (CO1), [CO-2], (Course Outcome 3)
    co_match = re.search(r'[\(\[\{]\s*(?:CO|C\d{3}\.?)\s*[-–:]?\s*([1-6]|CO[1-6])\s*[\)\]\}]', text, re.IGNORECASE)
    if co_match:
        val = co_match.group(1).upper()
        if val.startswith('CO'):
            co = val
        elif val.isdigit():
            co = f"CO{val}"
        text = text[:co_match.start()] + text[co_match.end():]
        
    # Extract combined (K2, CO1) or (CO1, K2)
    combo_match = re.search(r'[\(\[\{]\s*(?:(K[1-6]|\d)\s*,\s*(CO[1-6]|\d)|(CO[1-6]|\d)\s*,\s*(K[1-6]|\d))\s*[\)\]\}]', text, re.IGNORECASE)
    if combo_match:
        g = combo_match.groups()
        for item in g:
            if not item: continue
            item_u = item.upper()
            if 'K' in item_u or (item.isdigit() and not kl):
                kl = f"K{item_u.replace('K','')}"
            elif 'CO' in item_u:
                co = item_u
        text = text[:combo_match.start()] + text[combo_match.end():]
        
    # Extract marks e.g. [2 Marks], (13), (16)
    marks_match = re.search(r'[\(\[\{]\s*(\d{1,2})\s*(?:marks?|m)?\s*[\)\]\}]$', text, re.IGNORECASE)
    if marks_match:
        val = int(marks_match.group(1))
        if val in [1, 2, 3, 5, 6, 7, 8, 10, 12, 13, 14, 15, 16]:
            marks = val
            text = text[:marks_match.start()]
            
    # Clean multiple spaces and trailing punctuation artifacts
    clean_text = re.sub(r'\s+', ' ', text).strip()
    clean_text = re.sub(r'\s*[\(\[\{]\s*[\)\]\}]\s*$', '', clean_text).strip()
    return clean_text, kl, co, marks


def classify_table_columns(rows_data: List[List[str]]) -> Dict[str, int]:
    """
    Intelligently determines column roles (q_idx, kl_idx, co_idx, sno_idx, marks_idx)
    by analyzing both header labels and cell contents across all rows.
    """
    if not rows_data or not rows_data[0]:
        return {"q_idx": 0, "kl_idx": -1, "co_idx": -1, "sno_idx": -1, "marks_idx": -1}

    num_cols = max(len(r) for r in rows_data)
    if num_cols == 1:
        return {"q_idx": 0, "kl_idx": -1, "co_idx": -1, "sno_idx": -1, "marks_idx": -1}

    header_row = [c.strip().lower() for c in rows_data[0]]
    q_idx = -1
    kl_idx = -1
    co_idx = -1
    sno_idx = -1
    marks_idx = -1

    # 1. Header Name Inspection
    for idx, c_text in enumerate(header_row):
        if re.match(r'^(s|sl|q|item)[\.\s]*no[\.]?$', c_text) or c_text in ['sno', 'qno', 'no', '#']:
            sno_idx = idx
        elif any(w in c_text for w in ['question', 'description', 'particulars', 'item description', 'q.text']):
            q_idx = idx
        elif any(w in c_text for w in ['knowledge', 'bloom', 'kl', 'k.l', 'level', 'bt level', 'bt']):
            kl_idx = idx
        elif any(w in c_text for w in ['course outcome', 'outcome', 'co', 'c.o']):
            co_idx = idx
        elif any(w in c_text for w in ['marks', 'mark', 'max mark']):
            marks_idx = idx

    # 2. Content-Aware Inspection across data rows
    data_rows = rows_data[1:] if len(rows_data) > 1 else rows_data
    if data_rows:
        col_avg_lengths = [0.0] * num_cols
        col_kl_scores = [0] * num_cols
        col_co_scores = [0] * num_cols
        col_sno_scores = [0] * num_cols
        col_marks_scores = [0] * num_cols

        for row in data_rows:
            for c_i in range(min(num_cols, len(row))):
                val = row[c_i].strip()
                col_avg_lengths[c_i] += len(val)
                val_u = val.upper()
                if re.match(r'^\d+[a-z]?$', val_u) or re.match(r'^(?:Q|S)?\d+$', val_u):
                    col_sno_scores[c_i] += 1
                if re.match(r'^(K[1-6]|REMEMBER|UNDERSTAND|APPLY|ANALYZE|EVALUATE|CREATE)$', val_u):
                    col_kl_scores[c_i] += 1
                if re.match(r'^(CO[1-6]|C\d{3}\.?\d?)$', val_u):
                    col_co_scores[c_i] += 1
                if val.isdigit() and int(val) in [1, 2, 3, 5, 8, 10, 12, 13, 14, 15, 16]:
                    col_marks_scores[c_i] += 1

        total_d_rows = max(len(data_rows), 1)
        col_avg_lengths = [l / total_d_rows for l in col_avg_lengths]

        # Best question column: longest text length
        if q_idx == -1 or col_avg_lengths[q_idx] < 10:
            best_len_col = max(range(num_cols), key=lambda i: col_avg_lengths[i])
            if col_avg_lengths[best_len_col] > 10:
                q_idx = best_len_col

        # Best KL column
        if kl_idx == -1:
            best_kl_col = max(range(num_cols), key=lambda i: col_kl_scores[i])
            if col_kl_scores[best_kl_col] > 0 and best_kl_col != q_idx:
                kl_idx = best_kl_col

        # Best CO column
        if co_idx == -1:
            best_co_col = max(range(num_cols), key=lambda i: col_co_scores[i])
            if col_co_scores[best_co_col] > 0 and best_co_col not in [q_idx, kl_idx]:
                co_idx = best_co_col

        # Best S.No column
        if sno_idx == -1:
            best_sno_col = max(range(num_cols), key=lambda i: col_sno_scores[i])
            if col_sno_scores[best_sno_col] > 0 and best_sno_col not in [q_idx, kl_idx, co_idx]:
                sno_idx = best_sno_col

    # Final safe fallbacks
    if q_idx == -1:
        q_idx = 1 if num_cols > 1 else 0

    return {
        "q_idx": q_idx,
        "kl_idx": kl_idx,
        "co_idx": co_idx,
        "sno_idx": sno_idx,
        "marks_idx": marks_idx
    }


def clean_pdf_cell_text(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', text)
    text = re.sub(r'(\w+)-\s+(\w+)', r'\1\2', text)
    return re.sub(r'\s+', ' ', text).strip()


def merge_continuation_text(prev_text: str, new_text: str) -> str:
    if not prev_text:
        return new_text
    if not new_text:
        return prev_text
    if prev_text.endswith('-'):
        return prev_text[:-1] + new_text
    return prev_text + ' ' + new_text


def is_pdf_table_header_row(cells: List[str]) -> bool:
    if not cells:
        return False
    c0 = cells[0].strip().lower()
    if re.match(r'^(s|q|sl)[\.\s]*no[\.]?$|^question(s)?$', c0):
        return True
    if re.match(r'^\d+[a-z]?$', c0):
        return False
    row_str = " ".join(cells).lower()
    if ('question' in row_str or 'description' in row_str) and ('kl' in row_str or 'co' in row_str or 'level' in row_str or 'bloom' in row_str):
        return True
    return False


def is_question_starter_line(text: str) -> bool:
    """Checks if a text line appears to start a new question."""
    if not text or len(text) < 3:
        return False
    clean = text.strip()
    # Matches: "1. ", "1) ", "Q1. ", "14(a) ", "(i) ", "a. "
    if re.match(r'^(?:Q\.?\s*)?\d+[a-z\)\.]', clean, re.IGNORECASE):
        return True
    if re.match(r'^\([a-z\d]+\)', clean, re.IGNORECASE):
        return True
    # Action verbs
    first_word = clean.split()[0].lower().rstrip(':,.-')
    action_starters = [
        'what', 'why', 'how', 'explain', 'describe', 'define', 'give', 'list',
        'state', 'discuss', 'differentiate', 'distinguish', 'compare', 'name',
        'which', 'write', 'illustrate', 'prove', 'derive', 'show', 'mention',
        'outline', 'elaborate', 'briefly', 'summarize', 'critically', 'develop',
        'formulate', 'design', 'draw', 'analyze', 'analyse', 'evaluate',
        'compute', 'determine', 'calculate', 'construct', 'implement', 'apply',
        'classify', 'contrast', 'interpret', 'sketch', 'justify', 'clarify'
    ]
    if first_word in action_starters:
        return True
    return False


def parse_question_bank_docx(file_bytes: bytes, subject_code: str, semester: str) -> List[Dict[str, Any]]:
    doc = docx.Document(io.BytesIO(file_bytes))
    questions: List[Dict[str, Any]] = []

    current_unit = "Unit I"
    current_part = "A"
    current_marks = 2

    # First pass: check if document starts with a unit header
    for p in doc.paragraphs[:10]:
        u = parse_unit_from_text(p.text.strip())
        if u:
            current_unit = u
            break

    # Extract all elements in body (both paragraphs and tables)
    for element in doc.element.body:
        tag = element.tag.split('}')[-1]

        if tag == 'p':
            p = Paragraph(element, doc)
            text = p.text.strip()
            if not text:
                continue

            u = parse_unit_from_text(text)
            if u:
                current_unit = u

            p_part, p_marks = parse_part_marks_from_text(text)
            if p_part:
                current_part = p_part
                current_marks = p_marks

        elif tag == 'tbl':
            table = Table(element, doc)
            if not table.rows:
                continue

            # Extract raw cell strings from table
            raw_rows = []
            for row in table.rows:
                row_cells = [c.text.strip() for c in row.cells]
                raw_rows.append(row_cells)

            if not raw_rows:
                continue

            # Check if Row 0 or table header is a unit/part title banner
            first_row_str = " ".join([c for c in raw_rows[0] if c])
            u_r0 = parse_unit_from_text(first_row_str)
            if u_r0:
                current_unit = u_r0
            p_part_r0, p_marks_r0 = parse_part_marks_from_text(first_row_str)
            if p_part_r0:
                current_part = p_part_r0
                current_marks = p_marks_r0

            # Classify columns intelligently
            cols = classify_table_columns(raw_rows)
            q_idx = cols["q_idx"]
            kl_idx = cols["kl_idx"]
            co_idx = cols["co_idx"]
            marks_idx = cols["marks_idx"]

            for row_idx, cells in enumerate(raw_rows):
                row_str = " ".join([c for c in cells if c])
                
                # Check for inline unit/part changes inside table rows
                u_row = parse_unit_from_text(row_str)
                if u_row:
                    current_unit = u_row
                p_part_row, p_marks_row = parse_part_marks_from_text(row_str)
                if p_part_row:
                    current_part = p_part_row
                    current_marks = p_marks_row

                # Skip header row or title rows
                if row_idx == 0:
                    first_cell = cells[0].lower() if cells else ""
                    if any(w in first_cell for w in ['s.no', 'q.no', 'sno', 'qno', 'unit', 'part', 'sl']) or 'question' in row_str.lower():
                        continue

                if (u_row or p_part_row) and (len(cells) <= q_idx or len(cells[q_idx]) < 15):
                    continue

                if len(cells) <= q_idx:
                    continue

                raw_q_text = cells[q_idx].strip()
                if not raw_q_text or raw_q_text.lower().startswith("question") or raw_q_text.lower() in ["s.no", "q.no", "description"]:
                    continue

                # Extract KL / CO / Marks from dedicated columns
                raw_kl = cells[kl_idx].strip() if (kl_idx >= 0 and len(cells) > kl_idx) else ""
                raw_co = cells[co_idx].strip() if (co_idx >= 0 and len(cells) > co_idx) else ""
                row_marks = current_marks
                if marks_idx >= 0 and len(cells) > marks_idx and cells[marks_idx].isdigit():
                    row_marks = int(cells[marks_idx])

                # Extract inline metadata if present in question text
                clean_q_text, in_kl, in_co, in_marks = extract_inline_metadata(raw_q_text)
                final_kl = raw_kl or in_kl or ""
                final_co = raw_co or in_co or infer_co_from_unit(current_unit)
                final_marks = in_marks or row_marks or current_marks

                if len(clean_q_text) >= 5:
                    questions.append({
                        "subject_code": subject_code,
                        "semester": semester,
                        "text": clean_q_text,
                        "unit": current_unit,
                        "part": current_part,
                        "marks": final_marks,
                        "kl": final_kl,
                        "co": final_co
                    })

    # Strategy 2: If 0 questions were parsed from tables, run Paragraph-Based Extraction (for tableless Word files like UNIT III.docx)
    if not questions:
        logger.info("0 questions found in tables. Activating paragraph-based DOCX question extractor...")
        current_unit = "Unit I"
        current_part = "A"
        current_marks = 2

        # Check document top for unit
        for p in doc.paragraphs[:10]:
            u = parse_unit_from_text(p.text.strip())
            if u:
                current_unit = u
                break

        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue

            u = parse_unit_from_text(text)
            if u:
                current_unit = u
                continue

            p_part, p_marks = parse_part_marks_from_text(text)
            if p_part:
                current_part = p_part
                current_marks = p_marks
                continue

            # Ignore header/title like "UNIT III: DISTRIBUTED MUTEX" or short titles
            if parse_unit_from_text(text) or parse_part_marks_from_text(text)[0]:
                continue

            if is_question_starter_line(text) or not questions:
                clean_q, in_kl, in_co, in_m = extract_inline_metadata(text)
                final_kl = in_kl or ""
                final_co = in_co or infer_co_from_unit(current_unit)
                final_m = in_m or current_marks

                if len(clean_q) >= 5:
                    questions.append({
                        "subject_code": subject_code,
                        "semester": semester,
                        "text": clean_q,
                        "unit": current_unit,
                        "part": current_part,
                        "marks": final_m,
                        "kl": final_kl,
                        "co": final_co
                    })
            elif questions and len(text) > 2:
                # Merge continuation line
                prev_text = questions[-1]["text"]
                clean_text, in_kl, in_co, in_m = extract_inline_metadata(text)
                questions[-1]["text"] = merge_continuation_text(prev_text, clean_text)
                if in_kl and not questions[-1]["kl"]:
                    questions[-1]["kl"] = in_kl
                if in_co and not questions[-1]["co"]:
                    questions[-1]["co"] = in_co
                if in_m:
                    questions[-1]["marks"] = in_m

    logger.info(f"Parsed {len(questions)} questions from DOCX file.")
    return questions


def parse_question_bank_pdf(file_bytes: bytes, subject_code: str, semester: str) -> List[Dict[str, Any]]:
    questions: List[Dict[str, Any]] = []

    current_unit = "Unit I"
    current_part = "A"
    current_marks = 2

    def _process_pdf_pages(pdf_doc, table_settings=None):
        nonlocal current_unit, current_part, current_marks
        parsed_q = []

        for page in pdf_doc.pages:
            tables = page.find_tables(table_settings=table_settings) if table_settings else page.find_tables()
            text_lines = page.extract_text_lines() or []

            items = []
            for l in text_lines:
                items.append(('text', l['top'], l['text']))
            for t in (tables or []):
                items.append(('table', t.bbox[1], t))

            items.sort(key=lambda x: x[1])

            for item_type, top_y, item in items:
                if item_type == 'text':
                    text_str = item.strip()
                    u = parse_unit_from_text(text_str)
                    if u:
                        current_unit = u
                    p_part, p_marks = parse_part_marks_from_text(text_str)
                    if p_part:
                        current_part = p_part
                        current_marks = p_marks

                elif item_type == 'table':
                    table_obj = item
                    table_data = table_obj.extract()
                    if not table_data or not table_data[0]:
                        continue

                    # Clean table data
                    cleaned_rows = [[clean_pdf_cell_text(c) for c in row] for row in table_data]

                    # Check title row
                    row0_text = " ".join([c for c in cleaned_rows[0] if c])
                    u_r0 = parse_unit_from_text(row0_text)
                    if u_r0:
                        current_unit = u_r0
                    p_part_r0, p_marks_r0 = parse_part_marks_from_text(row0_text)
                    if p_part_r0:
                        current_part = p_part_r0
                        current_marks = p_marks_r0

                    cols = classify_table_columns(cleaned_rows)
                    q_idx = cols["q_idx"]
                    kl_idx = cols["kl_idx"]
                    co_idx = cols["co_idx"]
                    marks_idx = cols["marks_idx"]
                    sno_idx = cols["sno_idx"]

                    for row_idx, cells in enumerate(cleaned_rows):
                        row_text = " ".join([c for c in cells if c])
                        u_row = parse_unit_from_text(row_text)
                        if u_row:
                            current_unit = u_row
                        p_part_row, p_marks_row = parse_part_marks_from_text(row_text)
                        if p_part_row:
                            current_part = p_part_row
                            current_marks = p_marks_row

                        if row_idx == 0 and (is_pdf_table_header_row(cells) or u_row or p_part_row):
                            continue

                        if u_row or p_part_row:
                            continue

                        if len(cells) <= q_idx:
                            continue

                        sno_cell = cells[sno_idx] if (sno_idx >= 0 and len(cells) > sno_idx) else (cells[0] if cells else "")
                        question_text = cells[q_idx].strip()
                        raw_kl = cells[kl_idx].strip() if (kl_idx >= 0 and len(cells) > kl_idx) else ""
                        raw_co = cells[co_idx].strip() if (co_idx >= 0 and len(cells) > co_idx) else ""
                        row_marks = current_marks
                        if marks_idx >= 0 and len(cells) > marks_idx and cells[marks_idx].isdigit():
                            row_marks = int(cells[marks_idx])

                        if not question_text or question_text.lower().startswith("question") or question_text.lower() == "s. no":
                            continue

                        clean_q, in_kl, in_co, in_m = extract_inline_metadata(question_text)
                        final_kl = raw_kl or in_kl or ""
                        final_co = raw_co or in_co or infer_co_from_unit(current_unit)
                        final_m = in_m or row_marks or current_marks

                        is_new_q = bool(re.match(r'^\d+[a-z]?$', sno_cell.lower())) or is_question_starter_line(clean_q)

                        if not is_new_q and parsed_q:
                            parsed_q[-1]["text"] = merge_continuation_text(parsed_q[-1]["text"], clean_q)
                            if final_kl and not parsed_q[-1]["kl"]:
                                parsed_q[-1]["kl"] = final_kl
                            if final_co and not parsed_q[-1]["co"]:
                                parsed_q[-1]["co"] = final_co
                        elif len(clean_q) >= 5:
                            parsed_q.append({
                                "subject_code": subject_code,
                                "semester": semester,
                                "text": clean_q,
                                "unit": current_unit,
                                "part": current_part,
                                "marks": final_m,
                                "kl": final_kl,
                                "co": final_co
                            })
        return parsed_q

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        questions = _process_pdf_pages(pdf)
        
        # Fallback 1: Text-based table strategy
        if not questions:
            logger.info("Default table extraction returned 0 questions. Trying text-strategy table extraction...")
            questions = _process_pdf_pages(pdf, table_settings={"vertical_strategy": "text", "horizontal_strategy": "text"})

        # Fallback 2: Text-line extraction if PDF has no tables
        if not questions:
            logger.info("Table extraction returned 0 questions. Trying text-line question parser fallback...")
            c_unit = "Unit I"
            c_part = "A"
            c_marks = 2

            for page in pdf.pages:
                page_text = page.extract_text() or ""
                for line in page_text.split('\n'):
                    line_str = line.strip()
                    if not line_str:
                        continue
                    u = parse_unit_from_text(line_str)
                    if u:
                        c_unit = u
                        continue
                    p_p, p_m = parse_part_marks_from_text(line_str)
                    if p_p:
                        c_part = p_p
                        c_marks = p_m
                        continue

                    if is_question_starter_line(line_str) or not questions:
                        clean_q, in_kl, in_co, in_m = extract_inline_metadata(line_str)
                        final_kl = in_kl or ""
                        final_co = in_co or infer_co_from_unit(c_unit)
                        final_m = in_m or c_marks

                        if len(clean_q) >= 5:
                            questions.append({
                                "subject_code": subject_code,
                                "semester": semester,
                                "text": clean_q,
                                "unit": c_unit,
                                "part": c_part,
                                "marks": final_m,
                                "kl": final_kl,
                                "co": final_co
                            })
                    elif questions and not line_str.lower().startswith("page ") and not line_str.lower().startswith("question bank"):
                        clean_q, in_kl, in_co, in_m = extract_inline_metadata(line_str)
                        questions[-1]["text"] = merge_continuation_text(questions[-1]["text"], clean_q)

    logger.info(f"Parsed {len(questions)} questions from PDF file.")
    return questions


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = docx.Document(io.BytesIO(file_bytes))
    full_text = []
    for p in doc.paragraphs[:30]:
        if p.text.strip():
            full_text.append(p.text.strip())
            
    if doc.tables:
        for row in doc.tables[0].rows[:10]:
            row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
            if row_text:
                full_text.append(row_text)
                
    return "\n".join(full_text)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        if pdf.pages:
            return pdf.pages[0].extract_text() or ""
    return ""


ROMAN_TO_NUM = {'I': 1, 'II': 2, 'III': 3, 'IV': 4, 'V': 5, 'VI': 6, 'VII': 7, 'VIII': 8}
NUM_TO_ROMAN = {1: 'I', 2: 'II', 3: 'III', 4: 'IV', 5: 'V', 6: 'VI', 7: 'VII', 8: 'VIII'}
YEAR_ROMAN = {1: 'I', 2: 'II', 3: 'III', 4: 'IV'}

def derive_year_from_semester(sem_val: Optional[str]) -> str:
    if not sem_val:
        return "II"
    s = sem_val.strip().upper()
    if s in ROMAN_TO_NUM:
        sem_num = ROMAN_TO_NUM[s]
    elif s.isdigit():
        sem_num = int(s)
    else:
        sem_num = None
        for rom, num in ROMAN_TO_NUM.items():
            if re.search(rf'\b{rom}\b', s):
                sem_num = num
                break
        if not sem_num:
            digits = re.findall(r'\d+', s)
            if digits:
                sem_num = int(digits[0])

    if sem_num:
        year_num = (sem_num + 1) // 2
        return YEAR_ROMAN.get(year_num, "II")
    return "II"


def parse_text_metadata(text: str) -> dict:
    metadata = {
        "subject_code": None,
        "subject_name": None,
        "semester": None,
        "year": None,
        "degree": None,
        "branch": None,
        "degree_branch": None,
        "degree_branch_sem": None,
        "regulation": None
    }
    
    match_both = re.search(r'(?:Sub\.[ \t]*Code/Sub\.[ \t]*Name|Sub[ \t]*Code[ \t]*/[ \t]*Sub[ \t]*Name)[ \t]*:[ \t]*([A-Za-z0-9\-]+)[ \t]*/[ \t]*([^\n\r]+)', text, re.IGNORECASE)
    if match_both:
        metadata["subject_code"] = match_both.group(1).strip()
        metadata["subject_name"] = match_both.group(2).strip()
    else:
        match_code = re.search(r'Subject[ \t]*Code[ \t]*:[ \t]*([A-Za-z0-9\-]+)', text, re.IGNORECASE)
        if match_code:
            metadata["subject_code"] = match_code.group(1).strip()
            
        match_name = re.search(r'Subject[ \t]*:[ \t]*([^:\n\r\t]+)', text, re.IGNORECASE)
        if match_name:
            name_val = match_name.group(1).strip()
            name_val = re.split(r'\s{2,}', name_val)[0]
            metadata["subject_name"] = name_val.strip()

    match_deg_br = re.search(r'(?:Degree[ \t]*/[ \t]*Branch)[ \t]*:[ \t]*([^\n\r\t]+)', text, re.IGNORECASE)
    if match_deg_br:
        val = match_deg_br.group(1).strip()
        val = re.split(r'\s{2,}', val)[0]
        parts = [p.strip() for p in val.split('/') if p.strip()]
        if len(parts) >= 2:
            metadata["degree"] = parts[0]
            metadata["branch"] = parts[1]
            metadata["degree_branch"] = f"{parts[0]}/{parts[1]}"
        elif len(parts) == 1:
            metadata["degree_branch"] = parts[0]
            if any(w in parts[0].upper() for w in ["B.E", "BE", "B.TECH", "BTECH"]):
                metadata["degree"] = parts[0]
            else:
                metadata["branch"] = parts[0]
    else:
        match_deg_br_sem = re.search(r'(?:Degree[ \t]*/[ \t]*Branch[ \t]*/[ \t]*Sem)[ \t]*:[ \t]*([^\n\r\t]+)', text, re.IGNORECASE)
        if match_deg_br_sem:
            val = match_deg_br_sem.group(1).strip()
            parts = [p.strip() for p in val.split('/') if p.strip()]
            if len(parts) >= 2:
                metadata["degree"] = parts[0]
                metadata["branch"] = parts[1]
                metadata["degree_branch"] = f"{parts[0]}/{parts[1]}"
            if len(parts) >= 3:
                metadata["semester"] = parts[2]
        else:
            match_deg = re.search(r'Degree[ \t]*:[ \t]*([^\n\r\t]+)', text, re.IGNORECASE)
            match_br = re.search(r'Branch[ \t]*:[ \t]*([^\n\r\t]+)', text, re.IGNORECASE)
            if match_deg:
                metadata["degree"] = match_deg.group(1).strip().split()[0]
            if match_br:
                metadata["branch"] = match_br.group(1).strip().split()[0]
            if metadata["degree"] or metadata["branch"]:
                metadata["degree_branch"] = f"{metadata.get('degree') or 'B.E'}/{metadata.get('branch') or 'CSE'}"

    match_yr_sem = re.search(r'(?:Year[ \t]*/[ \t]*Sem|Year[ \t]*/[ \t]*Semester)[ \t]*:[ \t]*([^\n\r\t]+)', text, re.IGNORECASE)
    if match_yr_sem:
        val = match_yr_sem.group(1).strip()
        val = re.split(r'\s{2,}', val)[0]
        parts = [p.strip() for p in val.split('/') if p.strip()]
        if len(parts) >= 2:
            metadata["year"] = parts[0]
            metadata["semester"] = parts[1]
        elif len(parts) == 1:
            metadata["semester"] = parts[0]

    if not metadata["semester"]:
        match_sem_line = re.search(r'(?:Sem|Semester)[ \t]*:[ \t]*([^\n\r\t]+)', text, re.IGNORECASE)
        if match_sem_line:
            sem_str = match_sem_line.group(1).strip()
            sem_str = re.split(r'\s{2,}', sem_str)[0]
            if "/" in sem_str:
                sem_str = sem_str.split("/")[-1].strip()
            metadata["semester"] = sem_str

    if metadata["semester"]:
        sem_clean = metadata["semester"].strip().upper()
        if sem_clean in ROMAN_TO_NUM:
            metadata["semester"] = sem_clean
        elif sem_clean.isdigit() and int(sem_clean) in NUM_TO_ROMAN:
            metadata["semester"] = NUM_TO_ROMAN[int(sem_clean)]

    if metadata["semester"] and not metadata["year"]:
        metadata["year"] = derive_year_from_semester(metadata["semester"])
    elif not metadata["year"]:
        metadata["year"] = "II"

    if not metadata["degree"]:
        metadata["degree"] = "B.E"
    if not metadata["branch"]:
        metadata["branch"] = "CSE"
    if not metadata["degree_branch"]:
        metadata["degree_branch"] = f"{metadata['degree']}/{metadata['branch']}"

    if metadata["semester"]:
        metadata["degree_branch_sem"] = f"{metadata['degree_branch']} / {metadata['semester']}"
    else:
        metadata["degree_branch_sem"] = metadata['degree_branch']

    match_reg = re.search(r'Regulation[ \t]*:[ \t]*(\d{4})', text, re.IGNORECASE)
    if match_reg:
        metadata["regulation"] = match_reg.group(1).strip()
    else:
        match_reg_dash = re.search(r'(\d{4})[ \t]*-[ \t]*Regulation', text, re.IGNORECASE)
        if match_reg_dash:
            metadata["regulation"] = match_reg_dash.group(1).strip()
        else:
            match_reg_paren = re.search(r'\((\d{4})[ \t]*-[ \t]*REGULATION\)', text, re.IGNORECASE)
            if match_reg_paren:
                metadata["regulation"] = match_reg_paren.group(1).strip()
            else:
                match_reg_word = re.search(r'Regulation[ \t]*(\d{4})', text, re.IGNORECASE)
                if match_reg_word:
                    metadata["regulation"] = match_reg_word.group(1).strip()

    return metadata
